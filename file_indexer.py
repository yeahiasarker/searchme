import os
from typing import List, Dict, Optional
import numpy as np
import faiss
from pathlib import Path
import pickle
import hashlib
from sentence_transformers import SentenceTransformer
import magic  # for file type detection
import datetime
from mutagen import File as MusicFile  # for music metadata
import mimetypes
from PIL import Image  # for image metadata
import PyPDF2  # for PDF metadata
import docx  # for Word documents

class FileMetadata:
    """Stores and manages metadata for indexed files."""

    CONTENT_PREVIEW_LENGTH = 1000

    def __init__(self, file_path: str):
        """Initialize file metadata.
        
        Args:
            file_path: Path to the file
            
        Raises:
            ValueError: If file is an unsupported type
        """
        self.path = file_path
        self.name = os.path.basename(file_path)
        self.extension = os.path.splitext(self.name)[1].lower()
        
        if self.extension.lower() == '.svg':
            raise ValueError("SVG files are not supported")
            
        self._init_basic_metadata()
        self._init_specific_metadata()

    def _init_basic_metadata(self):
        """Initialize basic file metadata."""
        stat = os.stat(self.path)
        self.size = stat.st_size
        self.created_time = datetime.datetime.fromtimestamp(stat.st_ctime)
        self.modified_time = datetime.datetime.fromtimestamp(stat.st_mtime)
        self.mime_type = magic.from_file(self.path, mime=True)

    def _init_specific_metadata(self):
        """Initialize format-specific metadata fields."""
        self.artist = None
        self.title = None
        self.duration = None
        self.dimensions = None
        self.page_count = None
        self.content = None
        
        self._extract_specific_metadata()

    def _extract_specific_metadata(self):
        try:
            if self.mime_type.startswith('text/'):
                self._extract_text_content()
            elif self.mime_type == 'application/pdf':
                self._extract_pdf_metadata()
                self._extract_pdf_content()
            elif self.extension == '.docx':
                self._extract_docx_metadata()
                self._extract_docx_content()
            elif self.mime_type.startswith('audio/'):
                self._extract_audio_metadata()
            elif self.mime_type.startswith('image/'):
                self._extract_image_metadata()
        except Exception as e:
            print(f"Error extracting metadata/content for {self.path}: {str(e)}")

    def _extract_audio_metadata(self):
        try:
            audio = MusicFile(self.path)
            if audio:
                if hasattr(audio, 'tags') and audio.tags:
                    self.artist = str(audio.tags.get('artist', [None])[0])
                    self.title = str(audio.tags.get('title', [None])[0])
                if hasattr(audio, 'info') and hasattr(audio.info, 'length'):
                    self.duration = float(audio.info.length)
        finally:
            del audio  # Ensure the file handle is closed

    def _extract_image_metadata(self):
        with Image.open(self.path) as img:
            self.dimensions = tuple(img.size)  # Convert to regular tuple

    def _extract_pdf_metadata(self):
        with open(self.path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            self.page_count = len(pdf.pages)
            if pdf.metadata:
                # Convert potential PDF string objects to regular strings
                self.title = str(pdf.metadata.get('/Title', '')) if pdf.metadata.get('/Title') else None

    def _extract_docx_metadata(self):
        doc = docx.Document(self.path)
        self.page_count = len(doc.paragraphs)
        del doc  # Ensure the file handle is closed

    def _extract_text_content(self):
        try:
            with open(self.path, 'r', encoding='utf-8') as f:
                self.content = f.read()
        except UnicodeDecodeError:
            # Skip if file is not readable as text
            pass

    def _extract_pdf_content(self):
        with open(self.path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            content = []
            for page in pdf.pages:
                content.append(page.extract_text())
            self.content = '\n'.join(content)

    def _extract_docx_content(self):
        doc = docx.Document(self.path)
        self.content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        del doc

    def to_context_string(self) -> str:
        """Convert metadata and content to a searchable context string"""
        context_parts = [
            f"File name: {self.name}",
            f"Extension: {self.extension}",
            f"Size: {self.size} bytes",
            f"Created: {self.created_time.strftime('%Y-%m-%d %H:%M:%S')}",
            f"Modified: {self.modified_time.strftime('%Y-%m-%d %H:%M:%S')}",
            f"Type: {self.mime_type}"
        ]
        
        # Add specific metadata if available
        if self.artist:
            context_parts.append(f"Artist: {self.artist}")
        if self.title:
            context_parts.append(f"Title: {self.title}")
        if self.duration:
            context_parts.append(f"Duration: {self.duration:.2f} seconds")
        if self.dimensions:
            context_parts.append(f"Dimensions: {self.dimensions[0]}x{self.dimensions[1]}")
        if self.page_count:
            context_parts.append(f"Pages: {self.page_count}")
        
        # Add content if available, but truncate if too long
        if self.content:
            # Truncate content to a reasonable length (e.g., first 1000 characters)
            truncated_content = self.content[:1000].replace('\n', ' ').strip()
            if truncated_content:  # Only add if there's actual content
                context_parts.append(f"Content: {truncated_content}")
            
        return " | ".join(context_parts)

    def __getstate__(self):
        """Customize pickling behavior"""
        state = self.__dict__.copy()
        # Ensure all attributes are pickle-friendly
        for key, value in state.items():
            if isinstance(value, (tuple, list, dict, set)):
                state[key] = value.__class__(value)  # Create a new copy of containers
        return state

class FileIndexer:
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        self.encoder = SentenceTransformer(model_name)
        self.index = None
        self.file_mapping: Dict[int, str] = {}
        self.metadata_mapping: Dict[str, FileMetadata] = {}
        
        # Create a directory for storing index files
        self.index_dir = os.path.join(os.path.expanduser("~"), ".file_search_index")
        os.makedirs(self.index_dir, exist_ok=True)
        
        # Store index files in the index directory
        self.index_path = os.path.join(self.index_dir, "file_index.faiss")
        self.mapping_path = os.path.join(self.index_dir, "file_mapping.pkl")
        self.metadata_path = os.path.join(self.index_dir, "metadata_mapping.pkl")

    def create_file_embedding(self, file_path: str) -> np.ndarray:
        """Create embedding for a file"""
        try:
            metadata = FileMetadata(file_path)
            context = metadata.to_context_string()
            
            # Get embedding
            embedding = self.encoder.encode(context)
            return embedding.reshape(1, -1).astype('float32')
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return None

    def index_file(self, file_path: str) -> bool:
        """Index a single file"""
        try:
            if os.access(file_path, os.R_OK):
                embedding = self.create_file_embedding(file_path)
                
                if embedding is not None:
                    # Initialize index if not exists
                    if self.index is None:
                        dimension = self.encoder.get_sentence_embedding_dimension()
                        self.index = faiss.IndexFlatL2(dimension)
                    
                    file_count = len(self.file_mapping)
                    self.index.add(embedding)
                    self.file_mapping[file_count] = file_path
                    metadata = FileMetadata(file_path)
                    self.metadata_mapping[file_path] = metadata
                    return True
            return False
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return False

    def index_files(self, directory: str):
        """Index all files in a directory"""
        if not os.path.exists(directory):
            print(f"Directory not found: {directory}")
            return
        
        if os.path.isfile(directory):
            self.index_file(directory)
            return
        
        for root, _, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                self.index_file(file_path)

    def load_index(self) -> bool:
        """Load the index and mappings from disk"""
        try:
            if not all(os.path.exists(p) for p in [self.index_path, self.mapping_path, self.metadata_path]):
                print("Warning: One or more index files not found:")
                if not os.path.exists(self.index_path):
                    print(f"- Missing index file: {self.index_path}")
                if not os.path.exists(self.mapping_path):
                    print(f"- Missing mapping file: {self.mapping_path}")
                if not os.path.exists(self.metadata_path):
                    print(f"- Missing metadata file: {self.metadata_path}")
                return False

            self.index = faiss.read_index(self.index_path)
            with open(self.mapping_path, 'rb') as f:
                self.file_mapping = pickle.load(f)
            with open(self.metadata_path, 'rb') as f:
                self.metadata_mapping = pickle.load(f)
            
            print(f"Successfully loaded index with {len(self.file_mapping)} files")
            return True
            
        except Exception as e:
            print(f"Error loading index: {str(e)}")
            self.index = None
            self.file_mapping = {}
            self.metadata_mapping = {}
            return False

    def search(self, query: str, k: int = 5) -> List[Dict]:
        if self.index is None:
            raise ValueError("Index not created or loaded")
        
        query_vector = self.encoder.encode(query).reshape(1, -1).astype('float32')
        distances, indices = self.index.search(query_vector, k)
        
        results = []
        for i, idx in enumerate(indices[0]):
            if idx in self.file_mapping:
                file_path = self.file_mapping[idx]
                metadata = self.metadata_mapping.get(file_path)
                results.append({
                    'path': file_path,
                    'distance': float(distances[0][i]),
                    'metadata': metadata.__dict__ if metadata else None
                })
        
        return results 

    def save_index(self):
        """Save the index and mappings to disk"""
        try:
            if self.index is not None and len(self.file_mapping) > 0:
                print(f"Saving index to {self.index_dir}...")
                # Save index
                faiss.write_index(self.index, self.index_path)
                
                # Save file mapping
                with open(self.mapping_path, 'wb') as f:
                    pickle.dump(self.file_mapping, f)
                
                # Save metadata mapping
                with open(self.metadata_path, 'wb') as f:
                    pickle.dump(self.metadata_mapping, f)
                print(f"Successfully saved index with {len(self.file_mapping)} files")
                return True
            else:
                print("No files were indexed successfully")
                return False
        except Exception as e:
            print(f"Error saving index: {str(e)}")
            return False 